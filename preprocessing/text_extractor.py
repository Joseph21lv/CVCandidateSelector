#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Módulo para extraer texto de diferentes formatos de archivo (PDF, DOCX, TXT, imágenes).
"""

import os
import logging
from pathlib import Path
import pytesseract
from PIL import Image
import PyPDF2
import docx
import fitz  # PyMuPDF
import cv2
import numpy as np

logger = logging.getLogger(__name__)

class TextExtractor:
    """Clase para extraer texto de diferentes formatos de archivo."""
    
    def __init__(self, config):
        """
        Inicializa el extractor de texto.
        
        Args:
            config (Config): Objeto de configuración.
        """
        self.config = config
        # Configurar Tesseract para OCR
        self.tesseract_cmd = self.config.get('preprocessing', 'tesseract_cmd')
        if self.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd
    
    def extract(self, file_path):
        """
        Extrae texto de un archivo.
        
        Args:
            file_path (str or Path): Ruta al archivo.
            
        Returns:
            str: Texto extraído del archivo.
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"El archivo {file_path} no existe.")
            return ""
        
        # Determinar el tipo de archivo por su extensión
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif extension == '.docx':
                return self._extract_from_docx(file_path)
            elif extension == '.txt':
                return self._extract_from_txt(file_path)
            elif extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
                return self._extract_from_image(file_path)
            else:
                logger.warning(f"Formato de archivo no soportado: {extension}")
                return ""
        except Exception as e:
            logger.error(f"Error al extraer texto de {file_path}: {e}")
            return ""
    
    def _extract_from_pdf(self, file_path):
        """
        Extrae texto de un archivo PDF.
        
        Args:
            file_path (Path): Ruta al archivo PDF.
            
        Returns:
            str: Texto extraído del PDF.
        """
        try:
            # Intentar primero con PyMuPDF (más rápido y mejor con PDFs complejos)
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            
            doc.close()
            
            # Si no se extrajo texto, puede ser un PDF escaneado
            if not text.strip():
                return self._extract_from_scanned_pdf(file_path)
                
            return text
        except Exception as e:
            logger.warning(f"Error con PyMuPDF, intentando con PyPDF2: {e}")
            
            # Fallback a PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    
                    for page_num in range(len(reader.pages)):
                        text += reader.pages[page_num].extract_text() or ""
                
                # Si no se extrajo texto, puede ser un PDF escaneado
                if not text.strip():
                    return self._extract_from_scanned_pdf(file_path)
                    
                return text
            except Exception as e2:
                logger.error(f"Error al extraer texto con PyPDF2: {e2}")
                return ""
    
    def _extract_from_scanned_pdf(self, file_path):
        """
        Extrae texto de un PDF escaneado usando OCR.
        
        Args:
            file_path (Path): Ruta al archivo PDF.
            
        Returns:
            str: Texto extraído mediante OCR.
        """
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text += pytesseract.image_to_string(img, lang='spa+eng')
            
            doc.close()
            return text
        except Exception as e:
            logger.error(f"Error al extraer texto de PDF escaneado: {e}")
            return ""
    
    def _extract_from_docx(self, file_path):
        """
        Extrae texto de un archivo DOCX.
        
        Args:
            file_path (Path): Ruta al archivo DOCX.
            
        Returns:
            str: Texto extraído del DOCX.
        """
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error al extraer texto de DOCX: {e}")
            return ""
    
    def _extract_from_txt(self, file_path):
        """
        Extrae texto de un archivo TXT.
        
        Args:
            file_path (Path): Ruta al archivo TXT.
            
        Returns:
            str: Texto extraído del TXT.
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error al extraer texto de TXT: {e}")
            return ""
    
    def _extract_from_image(self, file_path):
        """
        Extrae texto de una imagen usando OCR.
        
        Args:
            file_path (Path): Ruta a la imagen.
            
        Returns:
            str: Texto extraído de la imagen.
        """
        try:
            # Cargar imagen con OpenCV para preprocesamiento
            img = cv2.imread(str(file_path))
            
            # Convertir a escala de grises
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Aplicar umbral adaptativo para mejorar el contraste
            thresh = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            
            # Reducir ruido
            kernel = np.ones((1, 1), np.uint8)
            img_erosion = cv2.erode(thresh, kernel, iterations=1)
            img_dilation = cv2.dilate(img_erosion, kernel, iterations=1)
            
            # Convertir a formato PIL para Tesseract
            img_pil = Image.fromarray(img_dilation)
            
            # Extraer texto con Tesseract
            text = pytesseract.image_to_string(img_pil, lang='spa+eng')
            
            return text
        except Exception as e:
            logger.error(f"Error al extraer texto de imagen: {e}")
            return ""
#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Módulo para extraer texto de diferentes formatos de documentos.
"""

import os
import logging
from pathlib import Path
import PyPDF2
import docx
import pytesseract
from PIL import Image
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

class TextExtractor:
    """Clase para extraer texto de diferentes formatos de documentos."""
    
    def __init__(self, config):
        """
        Inicializa el extractor de texto.
        
        Args:
            config (Config): Objeto de configuración.
        """
        self.config = config
    
    def extract(self, file_path):
        """
        Extrae texto de un archivo.
        
        Args:
            file_path (str or Path): Ruta al archivo.
            
        Returns:
            str: Texto extraído del archivo.
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"El archivo no existe: {file_path}")
            return ""
        
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif extension == '.txt':
                return self._extract_from_txt(file_path)
            elif extension in ['.jpg', '.jpeg', '.png']:
                return self._extract_from_image(file_path)
            else:
                logger.warning(f"Formato de archivo no soportado: {extension}")
                return ""
        except Exception as e:
            logger.error(f"Error al extraer texto de {file_path}: {e}")
            return ""
    
    def _extract_from_pdf(self, file_path):
        """
        Extrae texto de un archivo PDF.
        
        Args:
            file_path (Path): Ruta al archivo PDF.
            
        Returns:
            str: Texto extraído del PDF.
        """
        try:
            # Intentar primero con PyMuPDF (más rápido y mejor con PDFs complejos)
            doc = fitz.open(file_path)
            text = ""
            
            for page in doc:
                text += page.get_text()
            
            if text.strip():
                return text
            
            # Si PyMuPDF no extrae texto, intentar con PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page in reader.pages:
                    text += page.extract_text() or ""
            
            return text
        except Exception as e:
            logger.error(f"Error al extraer texto del PDF {file_path}: {e}")
            return ""
    
    def _extract_from_docx(self, file_path):
        """
        Extrae texto de un archivo DOCX.
        
        Args:
            file_path (Path): Ruta al archivo DOCX.
            
        Returns:
            str: Texto extraído del DOCX.
        """
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error al extraer texto del DOCX {file_path}: {e}")
            return ""
    
    def _extract_from_txt(self, file_path):
        """
        Extrae texto de un archivo TXT.
        
        Args:
            file_path (Path): Ruta al archivo TXT.
            
        Returns:
            str: Texto extraído del TXT.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Intentar con otra codificación
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error al extraer texto del TXT {file_path}: {e}")
                return ""
        except Exception as e:
            logger.error(f"Error al extraer texto del TXT {file_path}: {e}")
            return ""
    
    def _extract_from_image(self, file_path):
        """
        Extrae texto de una imagen usando OCR.
        
        Args:
            file_path (Path): Ruta a la imagen.
            
        Returns:
            str: Texto extraído de la imagen.
        """
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='spa')
            return text
        except Exception as e:
            logger.error(f"Error al extraer texto de la imagen {file_path}: {e}")
            return ""
